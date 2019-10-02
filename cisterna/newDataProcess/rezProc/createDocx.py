from docx import Document
from docx.shared import Inches
import os
import sys

args=sys.argv
docxname='test.docx' if len(args)==1 else args[1]

sosudPartsBot=['BOT-1', 'BOT-1-2', 'BOT-1-3', 'BOT-2', 'BOT-2-2','BOT-2-3']
sosudParts=['OBECH', 'UB+OB', 'UB-FAST', 'USIL']
zonesBot=[4,6,7,4,6,7]
zones=[2,1,3,5]
doc=Document()

views={1 : 'Слева', 2: 'Справа', 3: 'Вид снизу'}

ramaPic='ramaMises'
for v in [1,2]:
    doc.add_picture(ramaPic+'-{}.png'.format(v), width=Inches(6))
    doc.add_paragraph('Рисунок  - Эффективные напряжения по Мизесу. {}'.format(views[v]))

for v in [1,2]:
    doc.add_picture('emax-{}.png'.format(v), width=Inches(6))
    doc.add_paragraph('Рисунок  - Максимальная деформация растяжения в сосуде их ПКМ. {}'.format(views[v]))

for i, p in enumerate(sosudPartsBot):
    for v in [1,2,3]:
        fname='{}-{}.png'.format(p, v)
        if os.path.exists(fname):
            doc.add_picture(fname, width=Inches(6))
            doc.add_paragraph('Рисунок  - Критерий F. Область {}. {}'.format(zonesBot[i], views[v]))
            fname='gist-{}.png'.format(p)
            if os.path.exists(fname):
                doc.add_picture(fname, width=Inches(6))
                doc.add_paragraph('Рисунок  - Распределение критерия F по слоям. Область {}. {}'.format(zonesBot[i], views[v]))

for i, p in enumerate(sosudParts):
    for v in [1,2,3]:
        fname='{}-{}.png'.format(p, v)
        if os.path.exists(fname):
            doc.add_picture(fname, width=Inches(6))
            doc.add_paragraph('Рисунок  - Критерий F. Область {}. {}'.format(zones[i], views[v]))
    fname='gist-{}.png'.format(p)
    if os.path.exists(fname):
        doc.add_picture(fname, width=Inches(6))
        doc.add_paragraph('Рисунок  - Распределение критерия F по слоям. Область {}'.format(zones[i]))

doc.save(docxname)
