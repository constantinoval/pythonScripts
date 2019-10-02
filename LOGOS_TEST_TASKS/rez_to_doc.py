# -- coding: cp1251 --
from docx import Document
from numpy import linspace
from os.path import exists
from os import remove
from docx.shared import Cm


# task = {'name': 'tens11',
#         'description': 'Задача 1',
#         'eps': {'11': [0, 0.012, 0.012],
#                 '22': [0, 0, -0.015]
#                 }
#         }
# doc_path = '1.docx'

Npoints = len(list(task['eps'].values())[0])
t = linspace(0, 1, Npoints)

doc = Document(doc_path if exists(doc_path) else None)
doc.add_paragraph(task['description'])
doc.add_paragraph(
    'Заданные деформации элементарного объема')
tab = doc.add_table(rows=Npoints+1, cols=len(task['eps'])+1)
tab.rows[0].cells[0].text = 'время, мс'
for i, k in enumerate(task['eps'].keys()):
    tab.rows[0].cells[i+1].text = k
    for j in range(Npoints):
        tab.rows[1+j].cells[i+1].text = '{}'.format(task['eps'][k][j])
for j in range(Npoints):
    tab.rows[1+j].cells[0].text = '{:.2f}'.format(t[j])
doc.add_picture(task['name']+'-fig1.png', width=Cm(15.25))
doc.add_paragraph(
    'Заданные деформации элементарного объема')
doc.add_picture(task['name']+'-fig2.png', width=Cm(15.25))
doc.add_paragraph(
    'Сравнение компонент тензоров деформаций')
doc.add_picture(task['name']+'-fig3.png', width=Cm(15.25))
doc.add_paragraph(
    'История накопления поврежденности в элементарном объеме')
doc.save(doc_path)
for i in range(3):
    remove(task['name']+'-fig{}.png'.format(i+1))
