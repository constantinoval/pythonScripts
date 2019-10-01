import xlwings as xl
import matplotlib.pyplot as plt
from matplotlib import rcParams
rcParams['font.size'] = 14
import re
from collections import defaultdict
import numpy as np
import docx


def create_table(document, data):
    rows_count = len(data)
    cols_count = len(data[0])
    table = document.add_table(rows=rows_count, cols=cols_count)
    table.autofit = False
    table.style = document.styles['Table Grid']
    for i, r in enumerate(data):
        cells = table.rows[i].cells
        for j, c in enumerate(r):
            cells[j].text = c


def get_data(exp_code):
    found = False
    for sh in book.sheets:
        if sh.name == exp_code:
            found = True
            break
    if not found:
        return None
    sh = book.sheets[exp_code]
    data = sh.range('A2').options(transpose=True, expand='table').value
    return data


def format_data(data, fs='.2f'):
    if data == None:
        return '-'
    ffs = '{:' + fs + '}'
    return ffs.format(data)


def plot_pulses(t, ei, er, et, ej, fout='sync.png', sig=1.0, eps=None):
    summ = ei - er - et
    plt.figure(figsize=(10, 6))
    N = len(t)
    mev = N // 20
    if mev == 0:
        mev = 1
 #   mev = 0.1
    plt.plot(t, ei, 'k', marker='s', markevery=mev,
             markersize=6, label=u'падающий')
    plt.plot(t, er, 'k', marker='v', markevery=mev,
             markersize=6, label=u'отраженный')
    plt.plot(t, et, ':k', lw=2, label=u'прошедший')
    plt.plot(t, summ, '--k', label=u'сумма')
    plt.plot(t, ej, '--k', marker='o', markevery=mev,
             markersize=6, label=u'обойма')
    if eps:
        plt.axhline(eps, color='k', lw=1,
                    label=u'калибровка\nпо скорости\nударника')
    plt.axhline(0, color='k')
    plt.xlabel(u'время, мкс')
    plt.ylabel(u'деформация')
    plt.subplots_adjust(left=0.12, bottom=0.15, right=0.75)
    plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0.)
    plt.grid()
    plt.title(u'Импульсы деформации в мерных стержнях')
    plt.tight_layout()
    plt.savefig(fout)


def plot_diagr(data):
    f = plt.figure(figsize=(11, 5))
    ax1 = plt.subplot(121)
    ax1.plot(data[0], data[6] * 100, 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('осевая деф.,%')
    ax1 = plt.subplot(122)
    ax1.plot(data[0], data[5] * 100, 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('радиальная деф.,%')
    f.tight_layout()
    plt.savefig('tmp.png')

    f = plt.figure(figsize=(11, 5))
    ax1 = plt.subplot(121)
    ax1.plot(data[0], data[8], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('осевое напр.,МПа')
    ax1 = plt.subplot(122)
    ax1.plot(data[0], data[9], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('радиальное напр.,МПа')
    f.tight_layout()
    plt.savefig('tmp1.png')

    f = plt.figure(figsize=(11, 5))
    ax1 = plt.subplot(121)
    ax1.plot(data[0], data[10], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('давление,МПа')
    ax1 = plt.subplot(122)
    ax1.plot(data[0], data[11], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('интенсивность напр.,МПа')
    f.tight_layout()
    plt.savefig('tmp2.png')

    f = plt.figure(figsize=(11, 5))
    ax1 = plt.subplot(121)
    ax1.plot(data[7] * 100, data[10], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('объемная деф.,%')
    ax1.set_ylabel('давление,МПа')
    ax1 = plt.subplot(122)
    ax1.plot(data[6] * 100, data[8], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('осевая деф., %')
    ax1.set_ylabel('осевое напр.,МПа')
    f.tight_layout()
    plt.savefig('tmp3.png')


book = xl.Book('сжатие в обойме.xls')
exp_data = {}
sh = book.sheets['параметры испытаний']
data = sh.range('A5:S22').value
for d in data:
    exp_data[d[0]] = d
protokols = docx.Document()
ne = 1
for exp in exp_data.keys():
    print(exp)
    exp1 = get_data(exp.split('_')[-1])
    if not exp1:
        continue
    exp_par = exp_data[exp]
    protokols.add_heading(f'Протокол динамического испытания J-{ne:02d}')
    ne += 1
    p = protokols.add_paragraph()
    p.add_run('Параметры образца').bold = True
    data = [['материал', 'L0, мм', 'D0, мм', 'm, г',
             'плотность, кг/м^3']]
    idxs = [5, 6, 7, 8]
    data.append(['В25'] + [format_data(exp_par[i]) for i in idxs])
    create_table(protokols, data)
    p = protokols.add_paragraph()
    p.add_run('Параметры испытания').bold = True
    data = [['давление\n КВД,атм', 'скорость ударника, м/c',
             'ск.деф., 1/с']]
    idxs = [3, 4, 16]
    data.append([format_data(exp_par[i]) for i in idxs])
    create_table(protokols, data)
    p = protokols.add_paragraph()
    p.add_run('Примечание: ').bold = True
    p.add_run(exp_par[10])
    if exp1:
        exp1 = list(map(np.array, exp1))
        plot_pulses(exp1[0], -exp1[1], exp1[2], -exp1[3], exp1[4],
                    'tmp.png')
        protokols.add_picture('tmp.png', width=docx.shared.Inches(6))
        plot_diagr(exp1)
        protokols.add_picture('tmp.png', width=docx.shared.Inches(6))
        protokols.add_picture('tmp1.png', width=docx.shared.Inches(6))
        protokols.add_picture('tmp2.png', width=docx.shared.Inches(6))
        protokols.add_picture('tmp3.png', width=docx.shared.Inches(6))
        protokols.add_page_break()
        data = [['время, мкс', 'пад.\n*1000',
                 'отр.\n*1000', 'прош.\n*1000', 'обойма\n*1000']]
        NN = 45
        tt = np.linspace(0, max(exp1[0]), NN)
        newdata = [tt]
        for i, d in enumerate([1, 2, 3, 4]):
            newdata.append(np.interp(tt, exp1[0], exp1[1 + i]))
        for i in range(NN):
            table_data = [tt[i]]
            for j in [1, 2, 3, 4]:
                table_data.append(newdata[j][i] * 1000)
            data.append(list(map(lambda x: format_data(
                x, '.3f'), table_data)))
        create_table(protokols, data)
    protokols.add_page_break()
protokols.save('dyn_protocol-j.docx')
