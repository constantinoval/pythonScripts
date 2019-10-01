import xlwings as xl
import matplotlib.pyplot as plt
from matplotlib import rcParams
rcParams['font.size'] = 14
import re
from collections import defaultdict
import numpy as np
import docx
from math import log10


def format_num(num):
    if num == 0:
        return '0'
    n = log10(abs(num))
    fs = '{:.2f}'
    if 2 < n <= 4:
        fs = '{:.0f}'
    if -2 < n <= -1:
        fs = '{:.3f}'
    if n <= -2:
        fs = '{:.2e}'
    if 1 < n <= 2:
        fs = '{:.1f}'
    if n > 4:
        fs = '{:.2e}'
    return fs.format(num)


def create_table(document, data):
    rows_count = len(data)
    cols_count = len(data[0])
    table = document.add_table(rows=rows_count, cols=cols_count)
    table.autofit = False
    table.style = document.styles['Table Grid']
    for i, r in enumerate(data):
        cells = table.rows[i].cells
        for j, c in enumerate(r):
            cells[j].text = c


def get_data(exp_code):
    found = False
    for sh in book.sheets:
        if sh.name == exp_code:
            found = True
            break
    if not found:
        return None
    sh = book.sheets[exp_code]
    data = sh.range('A2').options(transpose=True, expand='table').value
    return data


def format_data(data, fs='.2f'):
    if data == None:
        return '-'
    #ffs = '{:' + fs + '}'
    return format_num(data)


def plot_diagr(data):
    f = plt.figure(figsize=(11, 5))
    ax1 = plt.subplot(121)
    ax1.plot(data[0], data[1], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('деформация')
    ax1 = plt.subplot(122)
    ax1.plot(data[0], data[2], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('напряжение, МПа')
    f.tight_layout()
    plt.savefig('tmp.png')

    f = plt.figure(figsize=(11, 5))
    ax1 = plt.subplot(121)
    ax1.plot(data[0], data[3], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('скорость деформации, 1/c')
    ax1 = plt.subplot(122)
    ax1.plot(data[1], data[2], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('дефорамация')
    ax1.set_ylabel('напряжение, МПа')
    ax2 = ax1.twinx()
    ax2.plot(data[1], data[3], '--k', lw=2)
    ax2.set_ylabel('скорость деформации, 1/c')
    ax2.set_ylim(ymax=3 * max(data[3]))
    f.tight_layout()
    plt.savefig('tmp2.png')
    plt.close('all')


book = xl.Book('сжатие.xls')
exp_data = {}
sh = book.sheets['параметры испытаний']
data = sh.range('A20:P33').value
for d in data:
    exp_data[d[0]] = d
protokols = docx.Document()
ne = 1
for exp in exp_data.keys():
    print(exp)
    exp1 = get_data(exp)
    if not exp1:
        continue
    exp_par = exp_data[exp]
    protokols.add_heading(f'Протокол динамического испытания C-{ne:02d}')
    p = protokols.add_paragraph()
    p.add_run(
        'Параметры установки: Классический метод Кольского. Сжатие.').bold = True
    create_table(protokols, [['', 'Длина, мм', 'Диаметр, мм', 'модуль упругости, ГПа', 'скорость звука, м/c'],
                             ['нагружающий стержень', '1500', '20', '70', '5150'],
                             ['опорный стержень', '1500', '20', '70', '5150'],
                             ['ударник', format_data(exp_par[1], '.0f'), '20', '70', '5150']])
    p = protokols.add_paragraph()
    p.add_run('Параметры образца').bold = True
    data = [['материал', 'L0, мм', 'D0, мм', 'm, г',
             'плотность, кг/м^3']]
    idxs = [5, 6, 7, 8]
    data.append(['В25'] + [format_data(exp_par[i]) for i in idxs])
    create_table(protokols, data)
    p = protokols.add_paragraph()
    p.add_run('Параметры испытания').bold = True
    data = [['давление КВД, атм', 'скорость ударника, м/c',
             'скорость деформации, 1/c', 'пред.деф.,%']]
    idxs = [3, 4, 12, 13]
    if exp_par[13]:
        exp_par[13] *= 100
    data.append([format_data(exp_par[i]) for i in idxs])
    create_table(protokols, data)
    exp1 = get_data(exp)
    if exp1:
        protokols.add_paragraph()
        plot_diagr(exp1)
        protokols.add_picture('tmp.png', width=docx.shared.Inches(6))
        protokols.add_picture('tmp2.png', width=docx.shared.Inches(6))
        data = [['время, мкс', 'деф.', 'напр.,\nМПа', 'ск.деф.,\n1/c',
                 'время, мкс', 'деф.', 'напр.,\nМПа', 'ск.деф.,\n1/c', ]]
        NN = 140
        tt = np.linspace(0, max(exp1[0]), NN)
        et = np.interp(tt, exp1[0], exp1[1])
        st = np.interp(tt, exp1[0], exp1[2])
        det = np.interp(tt, exp1[0], exp1[3])
        for i in range(NN // 2):
            data.append(list(map(lambda x: format_data(
                x, '.3f'), [tt[i], et[i], st[i], det[i],
                            tt[i + NN // 2], et[i + NN //
                                                2], st[i + NN // 2], det[i + NN // 2]
                            ])))
        create_table(protokols, data)
    protokols.add_page_break()
    ne += 1
    break
protokols.save('dyn_protocol-с.docx')
