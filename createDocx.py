import os
from collections import namedtuple, OrderedDict
from docx import Document
from docx.shared import Inches


def dataAnalisys():
    rez={'c': [], 't': []}
    i=namedtuple('expCondition', ['de', 'T'])
    for f in os.listdir(os.curdir):
        if os.path.splitext(f)[-1]=='.png':
            data=os.path.splitext(os.path.split(f)[-1])[0].split('-')
            if data[-1]=='all':
                rez[data[0][0]].append(i(data[1], data[2]))
    return rez

notes=OrderedDict({
    'all': 'Группа диаграмм (пунктир – скорость деформации).',
    'te': 'Изменение деформации образца во времени.',
    'ts': 'Изменение напряжения в образце во времени.',
    'tde': 'Изменение скорости деформации образца во времени.',
    'es': 'Средняя диаграмма.'
})
if __name__ == '__main__':
    s=dataAnalisys()
    for k in s:
        s[k].sort(key=lambda x: (x.T, float(x.de)))
        d=Document()
        for exp in s[k]:
            cond=f'T={exp.T} C, скорость деформации ~{exp.de} 1/c\n'
            for note in notes:
                w=6 if note=='all' else 4
                d.add_picture(f'{k}633-{exp.de}-{exp.T}-{note}.png', width=Inches(w))
                d.add_paragraph(f'Рисунок  – {notes[note]} \n'+cond)
        d.save(f'{k}.docx')