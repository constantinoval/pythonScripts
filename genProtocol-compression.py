# coding: cp1251
from math import log10
import docx
import numpy as np
from collections import defaultdict
import re
from odbcAccess import expODBC, calcDiagram2
import matplotlib.pyplot as plt
from matplotlib import rcParams
rcParams['font.size'] = 11


def format_num(num):
    if num == 0:
        return '0'
    n = log10(abs(num))
    fs = '{:.2f}'
    if 2 < n <= 4:
        fs = '{:.0f}'
    if -2 < n <= -1:
        fs = '{:.3f}'
    if n <= -2:
        fs = '{:.2e}'
    if 1 < n <= 2:
        fs = '{:.1f}'
    if n > 4:
        fs = '{:.2e}'
    return fs.format(num)


def create_table(document, data):
    rows_count = len(data)
    cols_count = len(data[0])
    table = document.add_table(rows=rows_count, cols=cols_count)
    table.autofit = False
    table.style = document.styles['Table Grid']
    for i, r in enumerate(data):
        cells = table.rows[i].cells
        for j, c in enumerate(r):
            cells[j].text = str(c)


def runningMeanFast(x, N):
    rez = np.roll(np.convolve(x, np.ones((N,))/N)[(N-1):], N//2)
    rez[:N//2] = rez[N//2]
    return rez


protokols = docx.Document()
db = expODBC(
    r"D:\experiments\База данных\db-work2.accdb")
exp_type = 'c'
mat_code = '661'
exp_type_name = 'сжатие'
notes_re = re.compile(r'(?P<w>[0-9.]+)\*(?P<h>[0-9.]+) *(?P<notes>[\w ]*)')


def get_line(e, s, n=50):
    n1 = s.argmax()
    fit = np.polyfit(e[n:n1], s[n:n1], deg=1)
    return fit, e[n:n1], np.poly1d(fit)(e[n:n1])


for num in db.getNumbers(exp_type, mat_code):
    exp_code = exp_type+mat_code+'-'+num['НомерОбразца']
    print(
        f'Формируется протокол испытания {exp_code}')
    try:
        experiment = db.getExperimentData(exp_code)
    except:
        continue
    protokols.add_heading(
        f'Протокол динамического испытания {exp_code}')
    protokols.add_paragraph(f'Дата испытания: {experiment.data}')
    protokols.add_paragraph(
        f'Параметры экспериментальной установки')
    b1 = db.getBarData(experiment.bars[0])
    table = [['Нагружающий стержень', '', '', ''],
             ['Материал', b1.mat, 'Длина, мм', b1.l],
             ['E, МПа', b1.E, 'Диаметр, мм', b1.d],
             ['c, м/c', b1.c, 'Тарир. коэффициент, 1/В',
                 experiment.tarir[0]],
             ['', '', 'Полож. датчиков, мм', experiment.datPosition[0]]
             ]
    b1 = db.getBarData(experiment.bars[1])
    table += [['Опорный стержень', '', '', ''],
              ['Материал', b1.mat, 'Длина, мм', b1.l],
              ['E, МПа', b1.E, 'Диаметр, мм', b1.d],
              ['c, м/c', b1.c, 'Тарир. коэффициент, 1/В',
               experiment.tarir[1]],
              ['', '', 'Полож. датчиков, мм', experiment.datPosition[1]]
              ]
    s = db.getStrickerData(experiment.striker)
    table += [['Ударник', '', '', ''],
              ['Материал', s.mat, 'Длина, мм', s.l],
              ['Диаметр, мм', s.d, '']
              ]
    create_table(protokols, table)

    protokols.add_paragraph('Параметры эксперимента')
    r = notes_re.search(experiment.note).groupdict()
    table = [['Тип эксперимента', exp_type_name, 'Давление в РК, атм', experiment.P],
             ['Скорость ударника, м/c',
                 experiment.V, 'T, С', experiment.T],
             ['Образец', '', '', ''],
             ['L, мм', experiment.l0, 'w, мм', r['w']],
             ['h, мм', r['h'], 'тарир.коэф., 1/В', experiment.tarir[2],
              ]
             ]
    create_table(protokols, table)
    if r['notes']:
        protokols.add_paragraph('Примечание: '+r['notes'])
    protokols.add_page_break()
    protokols.add_paragraph(
        'Осциллограммы (сверху вниз: нагружающий, опорный, датчик на образце)')
    f = plt.figure(figsize=(8, 8))
    experiment.osc['t'] *= 1e3
    for i in range(3):
        plt.subplot(f'31{i+1}')
        experiment.osc['rays'][i] = runningMeanFast(
            experiment.osc['rays'][i], 100)
        plt.plot(experiment.osc['t'], experiment.osc['rays'][i])
        plt.xlabel('время, мс')
        plt.ylabel('сигнал, В')
        plt.grid()
    f.tight_layout()
    plt.savefig('tmp.png')
    protokols.add_picture('tmp.png', width=docx.shared.Cm(15))
    plt.close('all')
    protokols.add_paragraph(
        f'Таблицы в файле: {exp_code}-osc.csv')
    np.savetxt(exp_code+'-osc.csv', np.array([experiment.osc['t']] +
                                             experiment.osc['rays']).T, delimiter=',',
               header='time_ms, input_bar_V, output_bar_V, sample_gauge_V')

    protokols.add_page_break()
    t, [et, st, det] = calcDiagram2(db, exp_code)
    t *= 1e3
    tt = experiment.osc['t']
    eg = -experiment.osc['rays'][2]*experiment.tarir[2]
    eg -= eg[:100].mean()
#    emax = eg[:100].max()*0.5
#    eg = runningMeanFast(eg, 200)
    emax = 1e-4
    for i, e in enumerate(eg):
        if e >= emax:
            break
    tt = tt[i:]
    tt -= tt[0]
    eg = eg[i:]
    deg = np.gradient(eg, tt*1e-3)

    f, ax = plt.subplots(4, 2, figsize=(8, 10))
    ax[0][0].plot(t, et)
    ax[0][0].grid(True)
    ax[0][0].set_xlabel('время, мс')
    ax[0][0].set_ylabel('деформация\n(по стержням)')

    ax[1][0].plot(tt, eg)
    ax[1][0].set_xlim(0, max(t))
    ax[1][0].set_xlabel('время, мс')
    ax[1][0].set_ylabel(
        'деформация\n(датчик на образце)')
    ax[1][0].grid(True)

    ax[2][0].plot(t, st)
    ax[2][0].set_xlabel('время, мс')
    ax[2][0].set_ylabel('напряжение, МПа')
    ax[2][0].grid(True)

    ax[0][1].plot(t, det)
    ax[0][1].set_xlabel('время, мс')
    ax[0][1].set_ylabel(
        'скорость деформации, 1/c\n(по стержням)')
    ax[0][1].grid(True)

    ax[1][1].plot(tt, deg)
    ax[1][1].grid(True)
    ax[1][1].set_xlim(0, max(t))
    ax[1][1].set_xlabel('время, мс')
    ax[1][1].set_ylabel(
        'скорость деформации, 1/c\n(по датчику на образце)')

    ax[2][1].plot(et, st)
    ax[2][1].set_xlabel(
        'деформация (по сигналам\nс мерных стержней)')
    ax[2][1].set_ylabel('напряжение, МПа')
    ax[2][1].grid(True)

    k, xx, yy = get_line(et, st)
    ax[2][1].plot(xx, yy, 'k')
    Ebars = k[0]

    axx = ax[2][1].twinx()
    axx.plot(et[::2], det[::2], 'r--')
    axx.set_ylabel(
        'скорость деформации, 1/c\n(по мерным стержням)\n - - - -')
    axx.set_ylim(top=axx.get_ylim()[1]*2)

    eg2 = np.interp(t, tt, eg)
    deg2 = np.interp(t, tt, deg)
    n1 = eg2.argmax()
    ax[3][0].plot(eg2[:n1], st[:n1], lw=2)
    ax[3][0].set_xlabel(
        'деформация (по датчику\nна образце)')
    ax[3][0].set_ylabel('напряжение, МПа')
    ax[3][0].grid(True)

    k, xx, yy = get_line(eg2[:n1], st[:n1])
    ax[3][0].plot(xx, yy, 'k')
    Egauge = k[0]
    Smax = st.max()
    de_bars = max(det)
    de_gauge = max(deg2)

    ax[3][1].axis('off')
    axx = ax[3][0].twinx()
    axx.plot(eg2[:n1:2], deg2[:n1:2], 'r--')
    axx.set_ylabel(
        'скорость деформации, 1/c\n(по датчику на образце)\n - - - -')
    axx.set_ylim(top=axx.get_ylim()[1]*2)
    f.tight_layout()
    plt.savefig('tmp.png')
    protokols.add_picture('tmp.png', width=docx.shared.Cm(15))
    protokols.add_paragraph(
        f'Таблицы в файле: {exp_code}-data.csv')
    data = [t]+experiment.pulses['pulses']+[et, st, det]
    data += [np.interp(t, tt, eg), np.interp(t, tt, deg)]
    np.savetxt(exp_code+'-data.csv', np.array(data).T,
               delimiter=',',
               header='time_ms, ei, er, et, eps, sig_MPa, de, e_gauge, de_gauge')
    protokols.add_page_break()

    protokols.add_paragraph('Результаты эксперимента')
    r = notes_re.search(experiment.note).groupdict()
    table = [['E(по мерным стержням), МПа', format_num(Ebars), 'E(по датчику на образце), МПа ', format_num(Egauge)],
             ['Скорость деформации(по мерным стержням), 1/c',
                 format_num(de_bars), 'Скорость деформации(по датчику на образце), 1/c', format_num(de_gauge)],
             ['Предел прочности, МПа',
                 format_num(Smax), '', '']
             ]
    create_table(protokols, table)
    protokols.add_page_break()
protokols.save('compression.docx')
