import xlwings as xl
import matplotlib.pyplot as plt
from matplotlib import rcParams
rcParams['font.size'] = 14
import re
from collections import defaultdict
import numpy as np
import docx


def create_table(document, data):
    rows_count = len(data)
    cols_count = len(data[0])
    table = document.add_table(rows=rows_count, cols=cols_count)
    table.autofit = False
    table.style = document.styles['Table Grid']
    for i, r in enumerate(data):
        cells = table.rows[i].cells
        for j, c in enumerate(r):
            cells[j].text = c


def get_data(exp_code):
    found = False
    for sh in book.sheets:
        if sh.name == exp_code:
            found = True
            break
    if not found:
        return None
    sh = book.sheets[exp_code]
    data = sh.range('A2').options(transpose=True, expand='table').value
    data2 = sh.range('F2').options(transpose=True, expand='table').value
    return data + data2


def format_data(data, fs='.2f'):
    if data == None:
        return '-'
    ffs = '{:' + fs + '}'
    return ffs.format(data)


def plot_pulses(t, ei, er, et, fout='sync.png', sig=1.0, eps=None):
    summ = ei - er - et
    plt.figure(figsize=(10, 6))
    N = len(t)
    mev = N // 20
    if mev == 0:
        mev = 1
 #   mev = 0.1
    plt.plot(t, ei, 'k', marker='s', markevery=mev,
             markersize=5, label=u'падающий')
    plt.plot(t, er, 'k', marker='v', markevery=mev,
             markersize=5, label=u'отраженный')
    plt.plot(t, et, ':k', lw=2, label=u'прошедший')
    plt.plot(t, summ, '--k', label=u'сумма')
    if eps:
        plt.axhline(eps, color='k', lw=1,
                    label=u'калибровка\nпо скорости\nударника')
    plt.axhline(0, color='k')
    plt.xlabel(u'время, мкс')
    plt.ylabel(u'деформация')
    plt.subplots_adjust(left=0.12, bottom=0.15, right=0.75)
    plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0.)
    plt.grid()
    plt.title(u'Импульсы деформации в мерных стержнях')
    plt.savefig(fout)


def plot_diagr(data):
    f = plt.figure(figsize=(11, 5))
    ax1 = plt.subplot(121)
    ax1.plot(data[0], data[-2] * 1e-3, 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('сила,кН')
    ax1 = plt.subplot(122)
    ax1.plot(data[0], data[-1], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('растяг.напряжение, МПа')
    f.tight_layout()
    plt.savefig('tmp.png')


book = xl.Book('раскалывание.xls')
exp_data = {}
sh = book.sheets['параметры испытаний']
data = sh.range('A35:M46').value
for d in data:
    exp_data[d[0]] = d
protokols = docx.Document()
ne = 1
for exp in exp_data.keys():
    print(exp)
    exp1 = get_data('t' + exp[-2:])
    if not exp1:
        continue
    exp_par = exp_data[exp]
    protokols.add_heading(f'Протокол динамического испытания T-{ne:02d}')
    ne += 1
    p = protokols.add_paragraph()
    p.add_run(
        'Параметры установки: Бразильский тест. Раскалывание.').bold = True
    create_table(protokols, [['', 'Длина, мм', 'Диаметр, мм', 'модуль упругости, ГПа', 'скорость звука, м/c'],
                             ['нагружающий стержень', '1500', '20', '70', '5150'],
                             ['опорный стержень', '1500', '20', '70', '5150'],
                             ['ударник', format_data(exp_par[1], '.0f'), '20', '70', '5150']])
    p = protokols.add_paragraph()
    p.add_run('Параметры образца').bold = True
    data = [['материал', 'L0, мм', 'D0, мм', 'm, г',
             'плотность, кг/м^3']]
    idxs = [5, 6, 7, 8]
    data.append(['В25'] + [format_data(exp_par[i]) for i in idxs])
    create_table(protokols, data)
    p = protokols.add_paragraph()
    p.add_run('Параметры испытания').bold = True
    data = [['скорость ударника, м/c', 'ск.роста напр., МПа/мкс', 'макс.напр.,МПа']]
    idxs = [4, 12, 11]
    data.append([format_data(exp_par[i]) for i in idxs])
    create_table(protokols, data)
    if exp1:
        exp1 = list(map(np.array, exp1))
        # protokols.add_paragraph()
        # plot_pulses(exp1[0], -exp1[1], exp1[2], -exp1[3],
        #             'tmp.png', eps=exp_par[4] / 2 / 5150)
        # protokols.add_picture('tmp.png', width=docx.shared.Inches(6))
        plot_diagr(exp1)
        protokols.add_picture('tmp.png', width=docx.shared.Inches(6))
        data = [['время, мкс', 'пад.\n*1000', 'отр.\n*1000', 'прош.\n*1000',
                 'сила,\nкН', 'раст.напр.\n,МПа']]
        NN = 56
        tt = np.linspace(0, max(exp1[0]), NN)
        newdata = [tt]
        for i, d in enumerate(exp1[1:]):
            newdata.append(np.interp(tt, exp1[0], exp1[1 + i]))
        for i in range(NN):
            table_data = [tt[i]]
            for j in [1, 2, 3]:
                table_data.append(newdata[j][i] * 1000)
            table_data.append(newdata[4][i] * 1e-3)
            table_data.append(newdata[5][i])
            data.append(list(map(lambda x: format_data(
                x, '.3f'), table_data)))
        create_table(protokols, data)
    protokols.add_page_break()
protokols.save('dyn_protocol-t.docx')
