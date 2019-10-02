import json
from math import sqrt
import docx

doc = docx.Document()

fasteners = json.load(open('fasteners.json'))
doc.add_paragraph('Расчет крепления бочки к раме')

r = 20e-3/2.
h = 30.2e-3
st=230e6
ssh=115e6
ssm_ub = 47.6e6
ssm_usil = 61e6

st_mean = fasteners['fn_mean']/3.14/r**2
ssh_mean = fasteners['fsh_mean']/3.14/r**2
ssm_mean = fasteners['fsh_mean']/2./r/h

K_st_mean=st/st_mean
K_ssh_mean=ssh/ssh_mean
K_ssm_mean_ub=ssm_ub/ssm_mean
K_ssm_mean_usil=ssm_usil/ssm_mean

table = doc.add_table(rows=2, cols=6)
table.rows[0].cells[0].text = 'нормальное усилие, Н'
table.rows[0].cells[1].text = 'перерезывающее усилие, Н'
table.rows[0].cells[2].text = 'коэффициент запаса на разрыв'
table.rows[0].cells[3].text = 'коэффициент запаса на срез'
table.rows[0].cells[4].text = 'коэффициент запаса на смятие (юбка)'
table.rows[0].cells[5].text = 'коэффициент запаса на смятие (оправка юбки)'

table.rows[1].cells[0].text = '{:.1f}'.format(fasteners['fn_mean'])
table.rows[1].cells[1].text = '{:.1f}'.format(fasteners['fsh_mean'])
table.rows[1].cells[2].text = '{:.1f}'.format(K_st_mean)
table.rows[1].cells[3].text = '{:.1f}'.format(K_ssh_mean)
table.rows[1].cells[4].text = '{:.1f}'.format(K_ssm_mean_ub)
table.rows[1].cells[5].text = '{:.1f}'.format(K_ssm_mean_usil)



bolts = json.load(open('bolts.json'))
def rb(Hname):
    N = int(Hname[1:])
    if  1<=N<=4:
        return 10e-3/2.
    else:
        return 22e-3/2.
def direction(Hmane):
    N = int(Hname[1:])
    if  1<=N<=4:
        return 1
    elif 5<=N<=8:
        return 2
    else:
        return 0

def sign(Hname):
    N = int(Hname[1:])
    if N in [7, 8]:
        return -1
    else:
      return 1

sb_n = 230e6
sb_sh = 115e6

doc.add_paragraph('Расчет болтовых соединений')
table = doc.add_table(rows=21, cols=5)
table.rows[0].cells[0].text = '№'
table.rows[0].cells[1].text = 'нормальное усилие, Н'
table.rows[0].cells[2].text = 'перерезывающее усилие, Н'
table.rows[0].cells[3].text = 'коэффициент запаса на разрыв'
table.rows[0].cells[4].text = 'коэффициент запаса на срез'
NR = 1
for Hname in sorted(bolts.keys(), key= lambda item: int(item[1:])):
    NN = direction(Hname)
    SS = [0,1,2]
    SS.pop(NN)
    ffn=sign(Hname)*bolts[Hname][NN]
    ffsh = sqrt(bolts[Hname][SS[0]]**2+bolts[Hname][SS[1]]**2)
    ssn = ffn/rb(Hname)**2/3.14
    sssh = ffsh/rb(Hname)**2/3.14
    Kn = sb_n/ssn if ffn>0 else 0
    Ksh = sb_sh/sssh
    table.rows[NR].cells[0].text = Hname
    table.rows[NR].cells[1].text = '{:.1f}'.format(ffn)
    table.rows[NR].cells[2].text = '{:.1f}'.format(ffsh)
    table.rows[NR].cells[3].text = '{:.1f}'.format(Kn)
    table.rows[NR].cells[4].text = '{:.1f}'.format(Ksh)

    NR+=1

doc.save('1.docx')