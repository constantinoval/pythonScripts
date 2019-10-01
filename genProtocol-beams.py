# coding: cp1251
from math import log10
import docx
import numpy as np
from collections import defaultdict
import re
from odbcAccess import expODBC, integrate
import matplotlib.pyplot as plt
from matplotlib import rcParams
rcParams['font.size'] = 11


def format_num(num):
    if num == 0:
        return '0'
    n = log10(abs(num))
    fs = '{:.2f}'
    if 2 < n <= 4:
        fs = '{:.0f}'
    if -2 < n <= -1:
        fs = '{:.3f}'
    if n <= -2:
        fs = '{:.2e}'
    if 1 < n <= 2:
        fs = '{:.1f}'
    if n > 4:
        fs = '{:.2e}'
    return fs.format(num)


def create_table(document, data):
    rows_count = len(data)
    cols_count = len(data[0])
    table = document.add_table(rows=rows_count, cols=cols_count)
    table.autofit = False
    table.style = document.styles['Table Grid']
    for i, r in enumerate(data):
        cells = table.rows[i].cells
        for j, c in enumerate(r):
            cells[j].text = str(c)


def runningMeanFast(x, N):
    rez = np.roll(np.convolve(x, np.ones((N,))/N)[(N-1):], N//2)
    rez[:N//2] = rez[N//2]
    return rez


def procBeam(db, exp_code):
    experiment = db.getExperimentData(exp_code)
    b1 = db.getBarData(experiment.bars[0])
    b2 = db.getBarData(experiment.bars[1])
    n = len(experiment.pulses['t'])
    dt = experiment.pulses['t'][1]-experiment.pulses['t'][0]
    U = []
    F = []
    V = []
    ein, eref, etr = experiment.pulses['pulses']
    for i in range(n):
        V.append(b1.c*(ein[i]+eref[i])-b2.c*etr[i])
        F.append(b2.E*b2.S*etr[i])
    U = integrate(V, dt)
    return U, V, F


protokols = docx.Document()
db = expODBC(
    r"D:\experiments\База данных\db-work2.accdb")
exp_type = 'b'
mat_code = '661'
exp_type_name = 'трехточечный изгиб балки'
notes_re = re.compile(
    r'(?P<l>[0-9.]+)\*(?P<w>[0-9.]+)\*(?P<h>[0-9.]+) *(?P<notes>[\w ]*)')


for num in db.getNumbers(exp_type, mat_code):
    exp_code = exp_type+mat_code+'-'+num['НомерОбразца']
    print(
        f'Формируется протокол испытания {exp_code}')
    try:
        experiment = db.getExperimentData(exp_code)
    except:
        continue
    protokols.add_heading(
        f'Протокол динамического испытания {exp_code}')
    protokols.add_paragraph(f'Дата испытания: {experiment.data}')
    protokols.add_paragraph(
        f'Параметры экспериментальной установки')
    b1 = db.getBarData(experiment.bars[0])
    table = [['Нагружающий стержень', '', '', ''],
             ['Материал', b1.mat, 'Длина, мм', b1.l],
             ['E, МПа', b1.E, 'Диаметр, мм', b1.d],
             ['c, м/c', b1.c, 'Тарир. коэффициент, 1/В',
              experiment.tarir[0]],
             ['', '', 'Полож. датчиков, мм', experiment.datPosition[0]]
             ]
    b1 = db.getBarData(experiment.bars[1])
    table += [['Опорный стержень', '', '', ''],
              ['Материал', b1.mat, 'Длина, мм', b1.l],
              ['E, МПа', b1.E, 'Диаметр, мм', b1.d],
              ['c, м/c', b1.c, 'Тарир. коэффициент, 1/В',
               experiment.tarir[1]],
              ['', '', 'Полож. датчиков, мм', experiment.datPosition[1]]
              ]
    s = db.getStrickerData(experiment.striker)
    table += [['Ударник', '', '', ''],
              ['Материал', s.mat, 'Длина, мм', s.l],
              ['Диаметр, мм', s.d, '', '']
              ]
    create_table(protokols, table)

    protokols.add_paragraph('Параметры эксперимента')
    r = notes_re.search(experiment.note).groupdict()
    table = [['Тип эксперимента', exp_type_name, 'Давление в РК, атм', experiment.P],
             ['Скорость ударника, м/c',
                 experiment.V, 'T, С', experiment.T],
             ['Образец', '', '', ''],
             ['L, мм', r['l'], 'w, мм', r['w']],
             ['h, мм', r['h'], '', ''],
             ]
    create_table(protokols, table)
    if r['notes']:
        protokols.add_paragraph('Примечание: '+r['notes'])
    protokols.add_page_break()
    protokols.add_paragraph(
        'Осциллограммы (сверху вниз: нагружающий, опорный)')
    f = plt.figure(figsize=(8, 8))
    experiment.osc['t'] *= 1e3
    for i in range(2):
        plt.subplot(f'31{i+1}')
        experiment.osc['rays'][i] = runningMeanFast(
            experiment.osc['rays'][i], 100)
        plt.plot(experiment.osc['t'], experiment.osc['rays'][i])
        plt.xlabel('время, мс')
        plt.ylabel('сигнал, В')
        plt.grid()
    f.tight_layout()
    plt.savefig('tmp.png')
    protokols.add_picture('tmp.png', width=docx.shared.Cm(15))
    plt.close('all')
    protokols.add_paragraph(
        f'Таблицы в файле: {exp_code}-osc.csv')
    np.savetxt(exp_code+'-osc.csv', np.array([experiment.osc['t']] +
                                             experiment.osc['rays']).T, delimiter=',',
               header='time_ms, input_bar_V, output_bar_V')

    protokols.add_page_break()
    t = experiment.pulses['t']
    U, V, F = procBeam(db, exp_code)
    t *= 1e3
    f, ax = plt.subplots(2, 2, figsize=(8, 6))
    ax[0][0].plot(t, np.array(U)*1000)
    ax[0][0].grid(True)
    ax[0][0].set_xlabel('время, мс')
    ax[0][0].set_ylabel('прогиб, мм')

    ax[0][1].plot(t, V)
    ax[0][1].grid(True)
    ax[0][1].set_xlabel('время, мс')
    ax[0][1].set_ylabel('скорость прогиба, м/c')

    ax[1][0].plot(t, np.array(F)*1e-3)
    ax[1][0].grid(True)
    ax[1][0].set_xlabel('время, мс')
    ax[1][0].set_ylabel('сила, кН')
    tau = 3./4.*np.array(F)/float(r['w'])/float(r['h'])
    ax[1][1].plot(t, tau)
    ax[1][1].grid(True)
    ax[1][1].set_xlabel('время, мс')
    ax[1][1].set_ylabel('сдвиговое напряжение, МПа')

    f.tight_layout()
    plt.savefig('tmp.png')
    protokols.add_picture('tmp.png', width=docx.shared.Cm(15))
    protokols.add_paragraph(
        f'Таблицы в файле: {exp_code}-data.csv')
    data = [t]+experiment.pulses['pulses']+[U, V, F, tau]
    np.savetxt(exp_code+'-data.csv', np.array(data).T,
               delimiter=',',
               header='time_ms, ei, er, et, U_m, V_m/s, F_N, tau_MPa')
    protokols.add_page_break()

protokols.save('beams.docx')
