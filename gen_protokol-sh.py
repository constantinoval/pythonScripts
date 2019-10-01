import xlwings as xl
import matplotlib.pyplot as plt
from matplotlib import rcParams
rcParams['font.size'] = 14
import re
from collections import defaultdict
import numpy as np
import docx
from math import log10


def create_table(document, data):
    rows_count = len(data)
    cols_count = len(data[0])
    table = document.add_table(rows=rows_count, cols=cols_count)
    table.autofit = False
    table.style = document.styles['Table Grid']
    for i, r in enumerate(data):
        cells = table.rows[i].cells
        for j, c in enumerate(r):
            cells[j].text = c


def format_num(num):
    if num == 0:
        return '0'
    n = log10(abs(num))
    fs = '{:.2f}'
    if 2 < n <= 4:
        fs = '{:.0f}'
    if -2 < n <= -1:
        fs = '{:.3f}'
    if n <= -2:
        fs = '{:.2e}'
    if 1 < n <= 2:
        fs = '{:.1f}'
    if n > 4:
        fs = '{:.2e}'
    return fs.format(num)


def get_data(exp_code):
    found = False
    for sh in book.sheets:
        if sh.name == exp_code:
            found = True
            break
    if not found:
        return None
    sh = book.sheets[exp_code]
    data = [sh.range('A2').options(transpose=True, expand='down').value]
    data.append(sh.range('G2').options(transpose=True, expand='down').value)
    return data


def format_data(data, fs='.2f'):
    if data == None:
        return '-'
    return format_num(data)


def plot_diagr(data):
    f = plt.figure(figsize=(6, 5))
    ax1 = plt.subplot(111)
    ax1.plot(data[0], data[1], 'k', lw=2)
    ax1.grid()
    ax1.set_xlabel('время, мкс')
    ax1.set_ylabel('напряжение, МПа')
    f.tight_layout()
    plt.savefig('tmp.png')
    plt.close('all')


book = xl.Book('shear.xls')
exp_data = {}
sh = book.sheets['Сжатие_20']
data = sh.range('A3:E4').value
for d in data:
    exp_data[d[0]] = d
protokols = docx.Document()
ne = 1
for exp in exp_data.keys():
    print(exp)
    exp1 = get_data(f'{int(exp)}-1')
    if not exp1:
        continue
    exp_par = exp_data[exp]
    protokols.add_heading(f'Протокол динамического испытания SH-{ne:02d}')
    ne += 1
    p = protokols.add_paragraph()
    p.add_run('Параметры испытания').bold = True
    data = [['давление КВД,атм', 'скорость ударника, м/c']]
    idxs = [3, 4]
    data.append([format_data(exp_par[i]) for i in idxs])
    create_table(protokols, data)
    if exp1:
        exp1 = list(map(np.array, exp1))
        plot_diagr(exp1)
        protokols.add_picture('tmp.png', width=docx.shared.Inches(4))
        data = [['время, мкс', 'напряжение, МПа',
                 'время, мкс', 'напряжение, МПа']]
        NN = 120
        tt = np.linspace(0, max(exp1[0]), NN)
        ss = np.interp(tt, exp1[0], exp1[1])
        for i in range(NN // 2):
            row1 = [format_num(tt[i]), format_num(ss[i]), format_num(
                tt[i] + NN // 2), format_num(ss[i] + NN // 2)]
            data.append(row1)
        create_table(protokols, data)
    protokols.add_page_break()
protokols.save('dyn_protocol-sh.docx')
